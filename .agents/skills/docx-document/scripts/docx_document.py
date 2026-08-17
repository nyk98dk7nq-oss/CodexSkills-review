#!/usr/bin/env python3
"""DOCXを位置情報付きMarkdownへ変換し、安全な限定編集を行う。"""

from __future__ import annotations

import argparse
from datetime import datetime
import json
import os
from pathlib import Path
import sys
from typing import Any, Iterator


def _require_runtime() -> None:
    if sys.version_info < (3, 12):
        raise RuntimeError(
            f"Python 3.12以上が必要です。現在: {sys.version.split()[0]}"
        )


def _require_docx() -> Any:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(
            "必要なライブラリ python-docx を読み込めません。READMEのインストール手順を確認してください。"
        ) from exc
    return docx


def _validate_input(path: Path) -> None:
    if not path.is_file():
        raise FileNotFoundError(f"入力ファイルが見つかりません: {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"DOCXファイルを指定してください: {path}")


def _ensure_distinct(input_path: Path, output_path: Path) -> None:
    if input_path.resolve() == output_path.resolve():
        raise ValueError("入力ファイルは上書きできません。別の出力パスを指定してください。")
    _validate_new_output_path(output_path, "出力先")


def _reject_symlink_components(path: Path, label: str) -> None:
    """出力パス自体と既存祖先のsymbolic linkを拒否する。"""
    absolute = path if path.is_absolute() else Path.cwd() / path
    for candidate in (absolute, *absolute.parents):
        if _is_link_like(candidate):
            raise ValueError(
                f"{label}にシンボリックリンクを使用できません: {candidate}"
            )


def _is_link_like(path: Path) -> bool:
    if path.is_symlink():
        return True
    is_junction = getattr(path, "is_junction", None)
    if is_junction is None:
        return False
    try:
        return bool(is_junction())
    except OSError:
        return False


def _validate_new_output_path(path: Path, label: str) -> None:
    _reject_symlink_components(path, label)
    if os.path.lexists(path):
        raise FileExistsError(
            f"{label}は既に存在します。上書きせず、何も書き込みません: {path}"
        )


def _validate_output_directory(path: Path, label: str) -> None:
    _reject_symlink_components(path, label)
    if os.path.lexists(path) and not path.is_dir():
        raise NotADirectoryError(f"{label}はディレクトリではありません: {path}")


def _repo_relative(path: Path, repo_root: Path) -> str:
    try:
        return path.resolve().relative_to(repo_root.resolve()).as_posix()
    except ValueError as exc:
        raise ValueError(
            f"入力ファイルはリポジトリルート配下に置いてください: {path}"
        ) from exc


def _yaml_string(value: str) -> str:
    return json.dumps(value, ensure_ascii=False)


def _md(value: Any) -> str:
    text = "" if value is None else str(value)
    return text.replace("\r\n", "\n").replace("\r", "\n").replace("|", "\\|").replace("\n", "<br>")


def _iter_blocks(parent: Any) -> Iterator[Any]:
    from docx.document import Document as DocumentType
    from docx.oxml.ns import qn
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    if isinstance(parent, DocumentType):
        parent_element = parent.element.body
    elif isinstance(parent, _Cell):
        parent_element = parent._tc
    else:
        raise TypeError("対応していないDOCXコンテナです。")

    for child in parent_element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _image_extension(part: Any) -> str:
    suffix = Path(str(part.partname)).suffix.lower()
    if suffix:
        return suffix
    return {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/gif": ".gif",
        "image/tiff": ".tiff",
        "image/bmp": ".bmp",
    }.get(getattr(part, "content_type", ""), ".bin")


def _paragraph_image_relationships(paragraph: Any) -> list[str]:
    from docx.oxml.ns import qn

    relationships: list[str] = []
    for run in paragraph.runs:
        for blip in run._element.xpath(".//a:blip"):
            relationship_id = blip.get(qn("r:embed"))
            if relationship_id:
                relationships.append(relationship_id)
    return relationships


def _render_order_paragraphs(document: Any) -> Iterator[Any]:
    for block in _iter_blocks(document):
        if hasattr(block, "rows"):
            for row in block.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
        else:
            yield block
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _plan_images(
    document: Any, images_dir: Path | None, source_stem: str
) -> list[dict[str, Any]]:
    """全画像の出力を先に確定し、衝突時は書込み前に停止する。"""
    plans: list[dict[str, Any]] = []
    planned_paths: set[str] = set()
    conflicts: list[Path] = []
    image_number = 0
    for paragraph in _render_order_paragraphs(document):
        for relationship_id in _paragraph_image_relationships(paragraph):
            related = paragraph.part.related_parts.get(relationship_id)
            plan: dict[str, Any] = {
                "relationship_id": relationship_id,
                "number": None,
                "path": None,
                "data": None,
                "extractable": related is not None and hasattr(related, "blob"),
            }
            if plan["extractable"]:
                image_number += 1
                plan["number"] = image_number
                if images_dir is not None:
                    image_path = images_dir / (
                        f"{source_stem}-image-{image_number:04d}{_image_extension(related)}"
                    )
                    _validate_new_output_path(image_path, "抽出画像の出力先")
                    path_key = str(image_path.resolve()).casefold()
                    if path_key in planned_paths:
                        conflicts.append(image_path)
                    planned_paths.add(path_key)
                    plan["path"] = image_path
                    plan["data"] = related.blob
            plans.append(plan)
    if conflicts:
        joined = "\n- ".join(str(path) for path in conflicts)
        raise FileExistsError(f"抽出画像の出力先が重複または既存です。何も書き込みません:\n- {joined}")
    return plans


def _paragraph_plain_text(paragraph: Any) -> str:
    iterator = getattr(paragraph, "iter_inner_content", None)
    if iterator is None:
        return paragraph.text
    return "".join(str(getattr(item, "text", "")) for item in iterator())


def _link_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("[", "\\[").replace("]", "\\]")


def _link_url(value: str) -> str:
    return value.replace(" ", "%20").replace("(", "\\(").replace(")", "\\)")


def _paragraph_markdown_text(paragraph: Any) -> str:
    iterator = getattr(paragraph, "iter_inner_content", None)
    if iterator is None:
        return paragraph.text
    fragments: list[str] = []
    for item in iterator():
        text = str(getattr(item, "text", ""))
        if item.__class__.__name__ == "Hyperlink":
            url = str(getattr(item, "url", "") or "")
            fragments.append(f"[{_link_text(text)}]({_link_url(url)})" if url else text)
        else:
            fragments.append(text)
    return "".join(fragments)


def _numbering_format(paragraph: Any, num_id: int, level: int) -> str:
    from docx.oxml.ns import qn

    try:
        numbering = paragraph.part.numbering_part.element
    except (AttributeError, NotImplementedError):
        return ""
    abstract_id = ""
    for element in numbering.findall(qn("w:num")):
        if element.get(qn("w:numId")) == str(num_id):
            reference = element.find(qn("w:abstractNumId"))
            abstract_id = reference.get(qn("w:val"), "") if reference is not None else ""
            break
    if not abstract_id:
        return ""
    for abstract in numbering.findall(qn("w:abstractNum")):
        if abstract.get(qn("w:abstractNumId")) != abstract_id:
            continue
        selected = None
        for level_element in abstract.findall(qn("w:lvl")):
            if level_element.get(qn("w:ilvl")) == str(level):
                selected = level_element
                break
        if selected is None:
            return ""
        number_format = selected.find(qn("w:numFmt"))
        return number_format.get(qn("w:val"), "") if number_format is not None else ""
    return ""


def _list_information(paragraph: Any) -> dict[str, Any] | None:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    style_id = paragraph.style.style_id if paragraph.style is not None else ""
    style_key = f"{style_name} {style_id}".casefold()
    paragraph_properties = paragraph._p.pPr
    num_pr = paragraph_properties.numPr if paragraph_properties is not None else None
    if num_pr is None and paragraph.style is not None:
        style_properties = paragraph.style.element.pPr
        num_pr = style_properties.numPr if style_properties is not None else None

    num_id: int | None = None
    level = 0
    if num_pr is not None:
        if num_pr.numId is not None and num_pr.numId.val is not None:
            num_id = int(num_pr.numId.val)
        if num_pr.ilvl is not None and num_pr.ilvl.val is not None:
            level = int(num_pr.ilvl.val)
    if num_id is None and not any(value in style_key for value in ("list", "bullet", "number", "箇条書き", "番号")):
        return None

    numbering_format = _numbering_format(paragraph, num_id, level) if num_id is not None else ""
    if numbering_format == "bullet" or "bullet" in style_key or "箇条書き" in style_key:
        kind = "箇条書き"
    elif numbering_format or "number" in style_key or "番号" in style_key:
        kind = "番号付き"
    else:
        kind = "リスト（種別不明）"
    return {"kind": kind, "level": level + 1, "num_id": num_id, "format": numbering_format}


def _paragraph_content(
    paragraph: Any,
    image_plans: list[dict[str, Any]],
    image_cursor: list[int],
    output_path: Path,
) -> tuple[str, list[str]]:
    references: list[str] = []
    for relationship_id in _paragraph_image_relationships(paragraph):
        if image_cursor[0] >= len(image_plans):
            raise RuntimeError("DOCX画像計画と文書順が一致しません。")
        plan = image_plans[image_cursor[0]]
        image_cursor[0] += 1
        if plan["relationship_id"] != relationship_id:
            raise RuntimeError("DOCX画像計画とrelationship IDが一致しません。")
        if not plan["extractable"]:
            references.append(f"[埋め込み画像: {relationship_id}（抽出不可）]")
            continue
        image_path = plan["path"]
        if image_path is None:
            references.append(f"[埋め込み画像: {relationship_id}]")
        else:
            relative_link = Path(os.path.relpath(image_path, output_path.parent)).as_posix()
            references.append(f"![埋め込み画像 {plan['number']}]({relative_link})")
    return _paragraph_markdown_text(paragraph), references


def _table_markdown(
    table: Any,
    table_index: int,
    image_plans: list[dict[str, Any]],
    image_cursor: list[int],
    output_path: Path,
) -> list[str]:
    values: list[list[str]] = []
    for row_index, row in enumerate(table.rows, start=1):
        row_values: list[str] = []
        for column_index, cell in enumerate(row.cells, start=1):
            fragments: list[str] = []
            for paragraph in cell.paragraphs:
                text, images = _paragraph_content(
                    paragraph, image_plans, image_cursor, output_path
                )
                if text:
                    fragments.append(text)
                fragments.extend(images)
            location = f"表 {table_index} 行 {row_index} 列 {column_index}"
            row_values.append(f"[{location}] " + "<br>".join(_md(value) for value in fragments))
        values.append(row_values)

    if not values:
        return ["（空の表）"]
    width = max(len(row) for row in values)
    normalized = [row + [""] * (width - len(row)) for row in values]
    lines = [
        "| " + " | ".join(f"列 {index}" for index in range(1, width + 1)) + " |",
        "| " + " | ".join("---" for _ in range(width)) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in normalized)
    return lines


def to_markdown(
    input_path: Path,
    output_path: Path,
    role: str,
    repo_root: Path,
    images_dir: Path | None,
) -> None:
    docx = _require_docx()
    _validate_input(input_path)
    _ensure_distinct(input_path, output_path)
    source_path = _repo_relative(input_path, repo_root)
    if images_dir is not None:
        _validate_output_directory(images_dir, "画像出力ディレクトリ")
    document = docx.Document(input_path)
    image_plans = _plan_images(document, images_dir, input_path.stem)
    image_cursor = [0]
    lines = [
        "---",
        f"source_path: {_yaml_string(source_path)}",
        f"source_name: {_yaml_string(input_path.name)}",
        'source_format: "docx"',
        f"document_role: {_yaml_string(role)}",
        f"converted_at: {_yaml_string(datetime.now().astimezone().isoformat(timespec='seconds'))}",
        'converter_skill: "docx-document"',
        "---",
        "",
        f"# {_md(input_path.name)}",
        "",
        "## 本文",
        "",
    ]

    paragraph_index = 0
    table_index = 0
    for block in _iter_blocks(document):
        if hasattr(block, "rows"):
            table_index += 1
            lines.extend([f"### 表 {table_index}", "", f"- 位置: `本文/表 {table_index}`", ""])
            lines.extend(
                _table_markdown(
                    block, table_index, image_plans, image_cursor, output_path
                )
            )
            lines.append("")
            continue

        paragraph_index += 1
        text, image_references = _paragraph_content(
            block, image_plans, image_cursor, output_path
        )
        if not text and not image_references:
            continue
        style_name = block.style.name if block.style is not None else ""
        lines.extend(
            [
                f"### 段落 {paragraph_index}",
                "",
                f"- 位置: `本文/段落 {paragraph_index}`",
                f"- スタイル: {_md(style_name) or 'なし'}",
                "",
            ]
        )
        list_information = _list_information(block)
        if list_information is not None:
            lines.append(f"- リスト種別: {list_information['kind']}")
            lines.append(f"- リストレベル: {list_information['level']}")
            if list_information["num_id"] is not None:
                lines.append(f"- numId: `{list_information['num_id']}`")
            if list_information["format"]:
                lines.append(f"- 番号形式: `{list_information['format']}`")
            lines.append("")
        if text:
            lines.append(text)
        lines.extend(image_references)
        lines.append("")

    for section_index, section in enumerate(document.sections, start=1):
        for container_name, container in (("ヘッダー", section.header), ("フッター", section.footer)):
            emitted_heading = False
            for paragraph_position, paragraph in enumerate(container.paragraphs, start=1):
                text, image_references = _paragraph_content(
                    paragraph, image_plans, image_cursor, output_path
                )
                if not text and not image_references:
                    continue
                if not emitted_heading:
                    lines.extend([f"## セクション {section_index} {container_name}", ""])
                    emitted_heading = True
                lines.extend(
                    [
                        f"### {container_name}段落 {paragraph_position}",
                        "",
                        f"- 位置: `セクション {section_index}/{container_name}/段落 {paragraph_position}`",
                        "",
                    ]
                )
                if text:
                    lines.append(text)
                lines.extend(image_references)
                lines.append("")

    if image_cursor[0] != len(image_plans):
        raise RuntimeError("DOCX画像計画をすべて消費できませんでした。")
    for image_plan in image_plans:
        image_path = image_plan["path"]
        if image_path is None:
            continue
        image_path.parent.mkdir(parents=True, exist_ok=True)
        with image_path.open("xb") as stream:
            stream.write(image_plan["data"])
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("x", encoding="utf-8", newline="\n") as stream:
        stream.write("\n".join(lines).rstrip() + "\n")


def _load_operations(path: Path) -> list[dict[str, Any]]:
    if not path.is_file():
        raise FileNotFoundError(f"操作JSONが見つかりません: {path}")
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"操作JSONを読み取れません: {path}: {exc}") from exc
    operations = payload.get("operations") if isinstance(payload, dict) else payload
    if not isinstance(operations, list) or not all(isinstance(item, dict) for item in operations):
        raise ValueError("操作JSONはオブジェクトの配列、または operations 配列を持つオブジェクトにしてください。")
    return operations


def _all_paragraphs(document: Any) -> Iterator[Any]:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _has_hyperlink(paragraph: Any) -> bool:
    return bool(paragraph._p.xpath("./w:hyperlink"))


def _replace_paragraph(paragraph: Any, old: str, new: str, limit: int | None) -> int:
    current = paragraph.text
    if not old or old not in current or limit == 0:
        return 0
    occurrences = current.count(old)
    replacement_count = occurrences if limit is None else min(occurrences, limit)
    replaced = current.replace(old, new, replacement_count)
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replaced)
    return replacement_count


def edit(input_path: Path, output_path: Path, operations_path: Path) -> None:
    docx = _require_docx()
    _validate_input(input_path)
    _ensure_distinct(input_path, output_path)
    document = docx.Document(input_path)

    for index, operation in enumerate(_load_operations(operations_path), start=1):
        operation_name = operation.get("op")
        try:
            if operation_name == "replace_text":
                old = operation["old"]
                new = operation["new"]
                if not isinstance(old, str) or not old:
                    raise ValueError("old は空でない文字列にしてください。")
                if not isinstance(new, str):
                    raise ValueError("new は文字列にしてください。")
                count = operation.get("count")
                if count is not None and (not isinstance(count, int) or count < 1):
                    raise ValueError("count は1以上の整数にしてください。")
                paragraphs = list(_all_paragraphs(document))
                if any(
                    _has_hyperlink(paragraph) and old in _paragraph_plain_text(paragraph)
                    for paragraph in paragraphs
                ):
                    raise ValueError(
                        "replace_text はハイパーリンクを含む段落を変更できません。"
                        "リンクを保持するため、その段落はWordで編集してください。"
                    )
                remaining = count
                replacements = 0
                for paragraph in paragraphs:
                    changed = _replace_paragraph(paragraph, old, new, remaining)
                    replacements += changed
                    if remaining is not None:
                        remaining -= changed
                        if remaining == 0:
                            break
                if replacements == 0:
                    raise ValueError(f"置換対象が見つかりません: {old}")
            elif operation_name == "add_paragraph":
                text = operation["text"]
                if not isinstance(text, str):
                    raise ValueError("text は文字列にしてください。")
                paragraph = document.add_paragraph()
                paragraph.add_run(text)
                if "style" in operation:
                    paragraph.style = operation["style"]
            elif operation_name == "set_table_cell":
                table_index = operation["table"]
                row_index = operation["row"]
                column_index = operation["column"]
                text = operation["text"]
                if not all(isinstance(value, int) and value >= 1 for value in (table_index, row_index, column_index)):
                    raise ValueError("table、row、column は1以上の整数にしてください。")
                if not isinstance(text, str):
                    raise ValueError("text は文字列にしてください。")
                document.tables[table_index - 1].rows[row_index - 1].cells[column_index - 1].text = text
            else:
                raise ValueError(f"許可されていない操作です: {operation_name}")
        except (IndexError, KeyError, TypeError, ValueError) as exc:
            raise ValueError(f"操作 {index} が不正です: {exc}") from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    subparsers = parser.add_subparsers(dest="command", required=True)
    markdown_parser = subparsers.add_parser("to-markdown", help="DOCXをMarkdownへ変換する")
    markdown_parser.add_argument("input", type=Path)
    markdown_parser.add_argument("output", type=Path)
    markdown_parser.add_argument("--role", required=True, choices=("checklist", "reference", "target"))
    markdown_parser.add_argument("--repo-root", required=True, type=Path)
    markdown_parser.add_argument("--images-dir", type=Path)

    edit_parser = subparsers.add_parser("edit", help="許可した操作だけでDOCXを編集する")
    edit_parser.add_argument("input", type=Path)
    edit_parser.add_argument("output", type=Path)
    edit_parser.add_argument("--operations", required=True, type=Path)
    return parser


def main() -> int:
    try:
        _require_runtime()
        arguments = build_parser().parse_args()
        if arguments.command == "to-markdown":
            to_markdown(
                arguments.input,
                arguments.output,
                arguments.role,
                arguments.repo_root,
                arguments.images_dir,
            )
        else:
            edit(arguments.input, arguments.output, arguments.operations)
        return 0
    except Exception as exc:
        print(f"エラー: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
