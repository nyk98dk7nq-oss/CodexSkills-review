from __future__ import annotations

import json
from pathlib import Path
import subprocess
import sys
import tempfile
import unittest

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image


SCRIPT = Path(__file__).resolve().parents[1] / "scripts" / "docx_document.py"


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class DocxDocumentCliTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary.name)
        target_dir = self.root / "input" / "targets"
        target_dir.mkdir(parents=True)
        image_path = target_dir / "sample.png"
        Image.new("RGB", (24, 16), "blue").save(image_path)
        self.source = target_dir / "design.docx"
        document = Document()
        document.add_heading("基本設計", level=1)
        paragraph = document.add_paragraph()
        paragraph.add_run("旧").bold = True
        paragraph.add_run("名称を使用する")
        link_paragraph = document.add_paragraph("参照: ")
        add_hyperlink(link_paragraph, "OpenAI公式", "https://openai.com/docs?a=1")
        link_paragraph.add_run(" を確認する")
        document.add_paragraph("箇条書き項目", style="List Bullet")
        document.add_paragraph("番号付き項目", style="List Number")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "項目"
        table.cell(0, 1).text = "値"
        table.cell(1, 0).text = "方式"
        table.cell(1, 1).text = "同期"
        document.add_picture(str(image_path), width=Inches(0.25))
        document.add_picture(str(image_path), width=Inches(0.25))
        document.sections[0].header.paragraphs[0].text = "機密"
        document.save(self.source)

    def tearDown(self) -> None:
        self.temporary.cleanup()

    def run_cli(self, *arguments: object) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [sys.executable, str(SCRIPT), *(str(value) for value in arguments)],
            text=True,
            capture_output=True,
            check=False,
        )

    def make_symlink(
        self, link: Path, target: Path, *, target_is_directory: bool = False
    ) -> None:
        try:
            link.symlink_to(target, target_is_directory=target_is_directory)
        except (NotImplementedError, OSError) as exc:
            self.skipTest(f"シンボリックリンクを作成できません: {exc}")

    def test_to_markdown_extracts_body_table_header_and_image(self) -> None:
        output = self.root / "work" / "markdown" / "design.md"
        images = self.root / "work" / "images"
        result = self.run_cli(
            "to-markdown",
            self.source,
            output,
            "--role",
            "target",
            "--repo-root",
            self.root,
            "--images-dir",
            images,
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        text = output.read_text(encoding="utf-8")
        self.assertIn('source_path: "input/targets/design.docx"', text)
        self.assertIn('document_role: "target"', text)
        self.assertIn('converter_skill: "docx-document"', text)
        self.assertIn("基本設計", text)
        self.assertIn("表 1 行 1 列 1", text)
        self.assertIn("セクション 1 ヘッダー", text)
        self.assertIn("機密", text)
        self.assertIn("参照: [OpenAI公式](https://openai.com/docs?a=1) を確認する", text)
        self.assertIn("- リスト種別: 箇条書き", text)
        self.assertIn("- リスト種別: 番号付き", text)
        self.assertIn("![埋め込み画像", text)
        self.assertEqual(len(list(images.glob("design-image-*.png"))), 2)

    def test_image_conflict_stops_before_writing_any_output(self) -> None:
        output = self.root / "work" / "markdown" / "design.md"
        images = self.root / "input" / "targets"
        conflict = images / "design-image-0002.png"
        sentinel = "利用者の既存画像".encode("utf-8")
        conflict.write_bytes(sentinel)
        result = self.run_cli(
            "to-markdown",
            self.source,
            output,
            "--role",
            "target",
            "--repo-root",
            self.root,
            "--images-dir",
            images,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertEqual(conflict.read_bytes(), sentinel)
        self.assertFalse((images / "design-image-0001.png").exists())
        self.assertFalse(output.exists())

    def test_broken_planned_image_symlink_stops_all_output(self) -> None:
        output = self.root / "work" / "markdown" / "design.md"
        images = self.root / "work" / "images"
        images.mkdir(parents=True)
        broken = images / "design-image-0002.png"
        self.make_symlink(broken, self.root / "outside" / "missing.png")

        result = self.run_cli(
            "to-markdown",
            self.source,
            output,
            "--role",
            "target",
            "--repo-root",
            self.root,
            "--images-dir",
            images,
        )

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("シンボリックリンク", result.stderr)
        self.assertTrue(broken.is_symlink())
        self.assertFalse((images / "design-image-0001.png").exists())
        self.assertFalse(output.exists())

    def test_symlink_primary_output_parent_and_images_dir_are_rejected(self) -> None:
        work = self.root / "work"
        work.mkdir()
        broken_output = work / "broken.md"
        broken_target = self.root / "outside" / "missing.md"
        self.make_symlink(broken_output, broken_target)
        result = self.run_cli(
            "to-markdown",
            self.source,
            broken_output,
            "--role",
            "target",
            "--repo-root",
            self.root,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertTrue(broken_output.is_symlink())
        self.assertFalse(broken_target.exists())

        outside = self.root / "outside"
        outside.mkdir()
        output_parent_link = self.root / "edited-link"
        self.make_symlink(output_parent_link, outside, target_is_directory=True)
        operations = self.root / "empty-operations.json"
        operations.write_text("[]", encoding="utf-8")
        edited = output_parent_link / "edited.docx"
        result = self.run_cli("edit", self.source, edited, "--operations", operations)
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("シンボリックリンク", result.stderr)
        self.assertFalse((outside / "edited.docx").exists())

        images_link = self.root / "images-link"
        self.make_symlink(images_link, outside, target_is_directory=True)
        markdown = work / "safe.md"
        result = self.run_cli(
            "to-markdown",
            self.source,
            markdown,
            "--role",
            "target",
            "--repo-root",
            self.root,
            "--images-dir",
            images_link,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("シンボリックリンク", result.stderr)
        self.assertFalse(markdown.exists())
        self.assertEqual(list(outside.iterdir()), [])

    def test_edit_replaces_across_runs_adds_paragraph_and_sets_table_cell(self) -> None:
        operations = self.root / "operations.json"
        operations.write_text(
            json.dumps(
                {
                    "operations": [
                        {"op": "replace_text", "old": "旧名称", "new": "新名称"},
                        {"op": "set_table_cell", "table": 1, "row": 2, "column": 2, "text": "非同期"},
                        {"op": "add_paragraph", "text": "改訂済み"},
                    ]
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        output = self.root / "output" / "edited.docx"
        result = self.run_cli("edit", self.source, output, "--operations", operations)
        self.assertEqual(result.returncode, 0, result.stderr)
        edited = Document(output)
        self.assertIn("新名称を使用する", [paragraph.text for paragraph in edited.paragraphs])
        self.assertEqual(edited.tables[0].cell(1, 1).text, "非同期")
        self.assertEqual(edited.paragraphs[-1].text, "改訂済み")
        original = Document(self.source)
        self.assertIn("旧名称を使用する", [paragraph.text for paragraph in original.paragraphs])

    def test_edit_rejects_input_overwrite(self) -> None:
        operations = self.root / "operations.json"
        operations.write_text("[]", encoding="utf-8")
        result = self.run_cli("edit", self.source, self.source, "--operations", operations)
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("上書きできません", result.stderr)

    def test_edit_rejects_existing_output(self) -> None:
        operations = self.root / "operations.json"
        operations.write_text("[]", encoding="utf-8")
        output = self.root / "output" / "existing.docx"
        output.parent.mkdir(parents=True)
        sentinel = "保持".encode("utf-8")
        output.write_bytes(sentinel)
        result = self.run_cli("edit", self.source, output, "--operations", operations)
        self.assertNotEqual(result.returncode, 0)
        self.assertEqual(output.read_bytes(), sentinel)

        markdown = self.root / "work" / "existing.md"
        markdown.parent.mkdir(parents=True)
        markdown.write_text("保持", encoding="utf-8")
        converted = self.run_cli(
            "to-markdown",
            self.source,
            markdown,
            "--role",
            "target",
            "--repo-root",
            self.root,
        )
        self.assertNotEqual(converted.returncode, 0)
        self.assertEqual(markdown.read_text(encoding="utf-8"), "保持")

    def test_replace_text_rejects_hyperlink_paragraph_without_output(self) -> None:
        operations = self.root / "operations.json"
        operations.write_text(
            json.dumps(
                [{"op": "replace_text", "old": "OpenAI公式", "new": "公式文書"}],
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        output = self.root / "output" / "edited.docx"
        result = self.run_cli("edit", self.source, output, "--operations", operations)
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("ハイパーリンクを含む段落", result.stderr)
        self.assertFalse(output.exists())


if __name__ == "__main__":
    unittest.main()
