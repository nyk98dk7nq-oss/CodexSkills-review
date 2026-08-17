from __future__ import annotations

import importlib.util
import json
import tempfile
import unittest
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
ORCHESTRATE_SCRIPT = (
    ROOT
    / ".agents"
    / "skills"
    / "review-documents-orchestrator"
    / "scripts"
    / "orchestrate.py"
)
SPEC = importlib.util.spec_from_file_location("orchestrate_end_to_end", ORCHESTRATE_SCRIPT)
assert SPEC and SPEC.loader
orchestrate = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(orchestrate)

try:
    import openpyxl
    from docx import Document
except ImportError:  # pragma: no cover - 事前環境確認の対象
    openpyxl = None
    Document = None


@unittest.skipIf(openpyxl is None or Document is None, "openpyxl と python-docx が必要です")
class EndToEndReviewTest(unittest.TestCase):
    def test_prepare_review_and_finalize(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            repo = Path(directory)
            for relative in (
                "input/checklists",
                "input/references",
                "input/targets",
                "output/reviews",
                "output/edited",
                "work",
            ):
                (repo / relative).mkdir(parents=True)

            checklist = repo / "input" / "checklists" / "design-review.xlsx"
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "基本設計"
            sheet.append(["ID", "チェック項目"])
            sheet.append(["C-01", "目的が明記されている"])
            sheet.append(["C-02", "障害時の復旧手順が明記されている"])
            workbook.save(checklist)
            workbook.close()

            target = repo / "input" / "targets" / "basic-design.docx"
            document = Document()
            document.add_heading("システムの目的", level=1)
            document.add_paragraph("このシステムは申請処理を一元化する。")
            document.save(target)

            run_id = "202608171230"
            manifest_path = orchestrate.prepare_repository(repo, run_id)
            self.assertTrue(manifest_path.is_file())
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(".", manifest["root"])
            self.assertNotIn(str(repo), manifest_path.read_text(encoding="utf-8"))
            self.assertIn(f"work/markdown/{run_id}/", manifest["targets"][0]["markdown"])

            bundle_path = repo / "work" / "review-runs" / run_id / "review_bundle.json"
            results = json.loads(bundle_path.read_text(encoding="utf-8"))
            for index, item in enumerate(results["items"]):
                if index == 0:
                    item.update(
                        {
                            "result": "適合",
                            "comment": "basic-design.docx.md: システムの目的 で目的の記述を確認した。",
                            "evidence": ["basic-design.docx.md: システムの目的"],
                            "improvement": "",
                        }
                    )
                else:
                    item.update(
                        {
                            "result": "不適合",
                            "comment": "basic-design.docx.md: 全文確認 で復旧手順の記述が存在しない。",
                            "evidence": ["basic-design.docx.md: 全文確認"],
                            "improvement": "障害検知、連絡、切戻し、復旧確認の担当者と手順を追記する。",
                        }
                    )
            results_path = repo / "work" / "review-runs" / run_id / "results.json"
            results_path.write_text(
                json.dumps(results, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
            )

            output_directory = orchestrate.finalize_repository(repo, run_id, results_path)
            output_checklist = output_directory / checklist.name
            summary = output_directory / "summary.md"
            self.assertTrue(output_checklist.is_file())
            self.assertTrue(summary.is_file())

            original = openpyxl.load_workbook(checklist)
            reviewed = openpyxl.load_workbook(output_checklist)
            try:
                self.assertEqual(2, original["基本設計"].max_column)
                reviewed_sheet = reviewed["基本設計"]
                self.assertEqual("レビュー対象ファイル名", reviewed_sheet["C1"].value)
                self.assertEqual("basic-design.docx", reviewed_sheet["C2"].value)
                self.assertEqual("適合", reviewed_sheet["D2"].value)
                self.assertEqual("不適合", reviewed_sheet["D3"].value)
            finally:
                original.close()
                reviewed.close()

            summary_text = summary.read_text(encoding="utf-8")
            self.assertIn("## 不適合項目と改善案", summary_text)
            self.assertIn("障害検知、連絡、切戻し、復旧確認", summary_text)


if __name__ == "__main__":
    unittest.main()
